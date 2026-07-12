import os
import json
import base64
import urllib.request
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def download_mermaid_image(mermaid_code, filename):
    try:
        graphbytes = mermaid_code.encode("utf8")
        base64_bytes = base64.b64encode(graphbytes)
        base64_string = base64_bytes.decode("ascii")
        url = "https://mermaid.ink/img/" + base64_string
        
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response, open(filename, 'wb') as out_file:
            data = response.read()
            out_file.write(data)
        return True
    except Exception as e:
        print(f"Error descargando {filename}: {e}")
        return False

def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def set_apa_style(doc):
    # Set Margins (2.54 cm on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        
    # Set Normal Style (Times New Roman 12, Double space, First line indent)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.first_line_indent = Cm(1.27)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    
    # Configure Heading 1 (Level 1: Centered, Bold, Title Case)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = None # default black
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing = 2.0
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)
    h1.paragraph_format.keep_with_next = True

    # Configure Heading 2 (Level 2: Flush Left, Bold, Title Case)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = None
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = 2.0
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.keep_with_next = True

    # Configure Heading 3 (Level 3: Flush Left, Bold Italic, Title Case)
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = None
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.line_spacing = 2.0
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.keep_with_next = True
    
    # Add setting to force update fields on open
    settings = doc.settings.element
    updateFields = OxmlElement('w:updateFields')
    updateFields.set(qn('w:val'), 'true')
    settings.append(updateFields)

def main():
    doc = Document()
    set_apa_style(doc)

    # PORTADA (Sin sangría)
    p1 = doc.add_paragraph()
    p1.paragraph_format.first_line_indent = 0
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("Universidad Metropolitana de Asunción\nFacultad de Posgrado\nHabilitación Pedagógica\n")
    run1.bold = True
    
    try:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.first_line_indent = 0
        r_logo = p_logo.add_run()
        r_logo.add_picture('c:\\Users\\HP 250 G10\\Documents\\GITHUT\\MonografiaUMA\\logo.jpeg', width=Inches(3.0))
    except Exception as e:
        print("Logo not found or could not be loaded:", e)
        
    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = 0
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("\nTrabajo Final de Monografía\n\nModelo de Estilos de Aprendizaje VAK\n\nAutor: Christhian José Raúl Keim Jara\nTutor: XXXXXX\n\n\n\nAsunción – Paraguay\nAño 2025")
    run2.bold = True
    doc.add_page_break()

    # DEDICATORIA
    doc.add_heading('Dedicatoria', level=1)
    doc.add_paragraph("A mi familia, por apoyarme incondicionalmente y darme un ejemplo constante a seguir. A los profesores que supieron guiarme, siempre con paciencia pero sin dejar de exigirme. A todas las personas que, como yo, confían en que la educación tiene el poder de cambiar nuestras vidas.")
    doc.add_page_break()

    # AGRADECIMIENTOS
    doc.add_heading('Agradecimientos', level=1)
    doc.add_paragraph("Quisiera expresar mi gratitud a la Universidad Metropolitana de Asunción, que me dio las herramientas clave para llevar adelante este trabajo. A mi orientador, por acompañarme de cerca y aportarme observaciones muy valiosas en el camino. También a mis compañeros, por las charlas, las ideas compartidas y ese aliento que nunca faltó. Y, por supuesto, a mi familia, por entender mis tiempos y apoyarme en cada paso de este proceso.")
    doc.add_page_break()

    # RESUMEN
    doc.add_heading('Resumen', level=1)
    doc.add_paragraph("Este trabajo analiza el modelo de estilos de aprendizaje VAK (Visual, Auditivo y Kinestésico), buscando entender sus bases teóricas y su potencial en el aula. El objetivo fue identificar qué caracteriza a cada preferencia sensorial para proponer estrategias prácticas de enseñanza. A través de una revisión documental, se cruzaron opiniones a favor y críticas sobre la idea de encasillar a los alumnos. Se encontró que, si bien la ciencia actual cuestiona la existencia de estilos fijos, usar estímulos variados mediante enseñanza multimodal resulta de gran ayuda para motivar e incluir. En definitiva, el modelo VAK sirve como guía heurística para diversificar la enseñanza y no como etiqueta diagnóstica.")
    p_claves = doc.add_paragraph()
    p_claves.paragraph_format.first_line_indent = Cm(1.27)
    r_claves_b = p_claves.add_run("Palabras clave: ")
    r_claves_b.italic = True
    p_claves.add_run("estilos de aprendizaje, VAK, preferencias sensoriales, revisión documental, enseñanza multimodal.")
    doc.add_page_break()

    # ABSTRACT
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph("This study analyzes the VAK (Visual, Auditory, and Kinesthetic) learning styles model, aiming to understand its theoretical foundations and potential in the classroom. The objective was to pinpoint what characterizes each sensory preference to propose practical teaching strategies. Through a documentary review, supporting views and criticisms regarding pigeonholing students were weighed. It was found that, although modern science questions the existence of fixed styles, using a variety of stimuli via multimodal teaching is highly helpful for motivating and including students. Ultimately, the VAK model serves as a heuristic guide to diversify teaching rather than a diagnostic label.")
    p_keys = doc.add_paragraph()
    p_keys.paragraph_format.first_line_indent = Cm(1.27)
    r_keys_b = p_keys.add_run("Keywords: ")
    r_keys_b.italic = True
    p_keys.add_run("learning styles, VAK, sensory preferences, documentary review, multimodal teaching.")
    doc.add_page_break()

    # ÍNDICE AUTOMÁTICO
    doc.add_heading('Índice', level=1)
    # Hint para el usuario (opcional)
    p_hint = doc.add_paragraph("(Nota: Al abrir en Word, puede presionar botón derecho aquí -> Actualizar campos para cargar la tabla de contenidos)")
    p_hint.paragraph_format.first_line_indent = 0
    p_hint.style.font.size = Pt(10)
    p_hint.style.font.italic = True
    add_toc(doc)
    doc.add_page_break()

    # CAPÍTULO I
    doc.add_heading('Capítulo I: Introducción y Metodología', level=1)
    
    doc.add_heading('1.1 Planteamiento del Problema', level=2)
    doc.add_paragraph("Cualquier docente nota enseguida que cada grupo de estudiantes es un mundo distinto, y que cada persona tiene su propia forma de engancharse con los temas. Por mucho tiempo, la enseñanza tendió a ser bastante uniforme, lo que a veces excluye a quienes procesan la información diferente. El modelo VAK se hizo muy popular para entender estas diferencias, pero se suele usar sin mucha reflexión, generando debate científico. Nos preguntamos: ¿hasta qué punto el modelo VAK sigue siendo útil para planificar clases sin caer en viejos neuromitos?")
    
    doc.add_heading('1.2 Justificación', level=2)
    doc.add_paragraph("Investigar esto tiene sentido porque los profesores necesitan herramientas reales y fundamentadas. Conocer la teoría y las críticas actuales a los estilos de aprendizaje evita que se etiquete a un estudiante para siempre, y en cambio, promueve clases más variadas e integradoras que suman a la práctica docente.")
    
    doc.add_heading('1.3 Objetivos', level=2)
    doc.add_paragraph("Objetivo general: Analizar mediante revisión documental el modelo de estilos de aprendizaje VAK, sus bases teóricas y su uso escolar real.")
    doc.add_paragraph("Objetivos específicos: Describir cómo se manifiestan las preferencias sensoriales. Comparar el modelo VAK clásico con la ciencia cognitiva actual. Sugerir estrategias de enseñanza variadas y multisensoriales.")
    
    doc.add_heading('1.4 Metodología', level=2)
    doc.add_paragraph("Se realizó una revisión documental y bibliográfica. Se rastrearon artículos y libros de los últimos 20 años en Scopus, SciELO y Google Scholar usando conceptos clave (modelo VAK, VARK, neuroeducación, enseñanza multimodal). Se incluyeron textos que defienden la utilidad práctica y aquellos con críticas desde la psicología cognitiva, agrupando luego la información para sacar conclusiones útiles.")
    doc.add_page_break()

    # CAPÍTULO II
    doc.add_heading('Capítulo II: Marco Teórico', level=1)
    
    doc.add_heading('2.1 Aclarando Conceptos', level=2)
    doc.add_paragraph("Es fundamental separar tres ideas: los 'Estilos de aprendizaje' (modelos teóricos amplios sobre cómo se asimila la información), las 'Preferencias sensoriales' (la inclinación particular por un sentido para recibir información) y las 'Estrategias didácticas' (los métodos que el docente planifica para facilitar la clase).")
    
    doc.add_heading('2.2 Del famoso VAK al VARK', level=2)
    doc.add_paragraph("El modelo VAK (Visual, Auditivo, Kinestésico) se hizo clásico en los años 80 y 90. En 1987, Neil Fleming introdujo el VARK añadiendo 'Lector/Escritor' (Read/Write). Aunque este trabajo se centra en el trío VAK, es importante reconocer que la lectura/escritura es un desprendimiento especializado de lo visual, como se muestra en los anexos.")
    
    doc.add_heading('2.3 Lo que Discuten Hoy los Científicos', level=2)
    doc.add_paragraph("Las neurociencias y la psicología cognitiva advierten que creer que un alumno aprende mejor solo si se le enseña en su 'estilo preferido' es un neuromito (Pashler et al., 2008). Las evidencias muestran que el aprendizaje se potencia mediante un enfoque multimodal, es decir, presentando la información combinada por distintas vías.")
    doc.add_page_break()

    # CAPÍTULO III
    doc.add_heading('Capítulo III: Desarrollo y Análisis del Modelo VAK', level=1)
    
    doc.add_heading('3.1 ¿Cómo son las Distintas Preferencias Sensoriales?', level=2)
    doc.add_paragraph("Aun sabiendo que no debemos etiquetar, el modelo VAK funciona como un catálogo de recursos.")
    
    doc.add_heading('La Preferencia Visual', level=3)
    doc.add_paragraph("Quienes tienen esta inclinación captan mejor la información con imágenes, gráficos y visiones globales. Suelen ser ordenados visualmente. Se benefician de mapas mentales, infografías y el uso de marcadores de color.")
    
    doc.add_heading('La Preferencia Auditiva', level=3)
    doc.add_paragraph("Destacan en lo que escuchan y debaten. Se enganchan mediante charlas o explicaciones en voz alta. Retienen muy bien lo dicho. Se favorecen con grupos de discusión, lectura en voz alta o podcasts.")
    
    doc.add_heading('La Preferencia Kinestésica', level=3)
    doc.add_paragraph("El aprendizaje pasa por el cuerpo, las manos y el hacer real. Tienen mucha memoria muscular y requieren movimiento. Sirven los laboratorios, juegos de rol, maquetas y simulaciones físicas.")
    
    doc.add_heading('3.2 El Camino de la Enseñanza Multisensorial', level=2)
    doc.add_paragraph("La mejor aplicación del VAK es dar clases multimodales. Se trata de explicar un tema por distintas vías a la vez (por ejemplo, hablar sobre un proceso, mostrar un diagrama y hacer un experimento corto). Así se incluyen a todos y se generan más conexiones cerebrales duraderas.")
    doc.add_page_break()

    # CAPÍTULO IV
    doc.add_heading('Capítulo IV: Conclusiones', level=1)
    doc.add_paragraph("1. Hay que repensar los estilos: El modelo VAK no es rígido; las preferencias son fluidas y varían según la tarea y el contexto.")
    doc.add_paragraph("2. Es una gran ayuda heurística para el docente: Nos obliga a darnos cuenta si nuestras clases están siendo demasiado homogéneas (por ejemplo, puramente auditivas) y a sumar más variedad.")
    doc.add_paragraph("3. Mezclar es ganar: La integración multisensorial rinde más. Combinar lo visual, auditivo y kinestésico hace clases más inclusivas y logra aprendizajes más sólidos.")
    
    doc.add_heading('Limitaciones del Trabajo', level=2)
    doc.add_paragraph("Como monografía de revisión, queda pendiente el trabajo de campo empírico en aulas reales. Asimismo, frente al debate actual donde los científicos descartan la rigidez de los estilos como neuromito, debe tomarse el modelo VAK con prudencia, como herramienta de flexibilización pedagógica y no como excusa para etiquetar alumnos.")
    doc.add_page_break()

    # BIBLIOGRAFÍA
    doc.add_heading('Bibliografía', level=1)
    
    # APA Bibliography requires hanging indent
    def add_bib_entry(doc, text):
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)

    add_bib_entry(doc, "Dunn, R., & Dunn, K. (2010). Teaching students through their individual learning styles. Allyn & Bacon.")
    add_bib_entry(doc, "Fleming, N. D. (2015). Teaching and learning styles: VARK strategies. Christchurch: VARK Learn Limited.")
    add_bib_entry(doc, "Gardner, H. (2011). Frames of mind: The theory of multiple intelligences. New York: Basic Books.")
    add_bib_entry(doc, "Kolb, D. A. (2014). Experiential learning: Experience as the source of learning and development. Pearson Education.")
    add_bib_entry(doc, "Pashler, H., McDaniel, M., Rohrer, D., & Bjork, R. (2008). “Learning styles: Concepts and evidence”. Psychological Science in the Public Interest, 9(3), 105–119. https://doi.org/10.1111/j.1539-6053.2009.01038.x")
    add_bib_entry(doc, "Salas, E., & Cannon-Bowers, J. A. (2001). “The science of training: A decade of progress”. Annual Review of Psychology, 52, 471–499. https://doi.org/10.1146/annurev.psych.52.1.471")
    add_bib_entry(doc, "Smith, M. K. (2002). “Howard Gardner and multiple intelligences”. The Encyclopedia of Informal Education. https://infed.org/mobi/howard-gardner-multiple-intelligences/")
    add_bib_entry(doc, "Universidad Metropolitana de Asunción. (2023). Guía para la presentación oral y defensa de tesis. UMA Documentos Académicos.")
    
    doc.add_page_break()

    # ANEXOS (Tablas)
    doc.add_heading('Anexos', level=1)
    doc.add_paragraph("Acá agregamos algunas tablas basadas en el modelo VARK para notar la evolución hacia el Lector/Escritor (Fleming, 2015).")
    
    doc.add_heading('Matriz: Comparativo Integral de Comportamientos', level=2)
    table1 = doc.add_table(rows=11, cols=4)
    table1.style = 'Table Grid'
    # Header
    hdr_cells1 = table1.rows[0].cells
    hdr_cells1[0].text = 'Categoría'
    hdr_cells1[1].text = 'Visual'
    hdr_cells1[2].text = 'Auditivo'
    hdr_cells1[3].text = 'Kinestésico'
    for cell in hdr_cells1:
        cell.paragraphs[0].runs[0].bold = True
    
    # Rows data
    rows_data = [
        ('Conducta',
         '- Son bastante organizados y prolijos.\n- Suelen quedarse observando.\n- Se les nota mucho lo que sienten en la cara.',
         '- Piensan en voz alta.\n- Mueven un poquito los labios cuando leen.\n- Tienen facilidad para charlar.',
         '- Mueven mucho las manos al hablar.\n- Siempre andan necesitando moverse.\n- Su cuerpo habla por ellos.'),
        ('Aprendizaje',
         '- Cazan los temas rápido si hay un esquema.\n- Necesitan ver el mapa completo.\n- Si se dice en el aire, olvidan fácil.',
         '- Les sirve estudiar repitiendo en voz alta.\n- Si se olvidan de un punto en el medio, se pierden.',
         '- Entienden mejor cuando "hacen".\n- Necesitan sentirse involucrados físicamente.'),
        ('Lectura',
         '- Les gustan las descripciones; a veces se quedan con la mirada perdida imaginándose la escena.',
         '- Prefieren los diálogos y las obras de teatro; evitan las descripciones largas.',
         '- Disfrutan las historias de acción donde ocurren cosas físicas.'),
        ('Ortografía',
         '- Tienen menos faltas. "Ven" las palabras antes de escribirlas.',
         '- Suelen cometer más faltas porque escriben según el sonido (como "suena").',
         '- Cometen faltas; suelen escribir la palabra y comprobar si "les da buena espina" verla.'),
        ('Memoria',
         '- Recuerdan lo que ven (ej. las caras), pero les cuesta recordar los nombres.',
         '- Recuerdan lo que oyen (ej. los nombres), pero olvidan las caras.',
         '- Recuerdan lo que hicieron o la impresión general que algo les causó, no los detalles.'),
        ('Imaginación',
         '- Piensan en imágenes de forma muy detallada.',
         '- Piensan en sonidos y no recuerdan tantos detalles visuales.',
         '- Sus imágenes mentales son pocas y poco detalladas, ligadas al movimiento.'),
        ('Almacenamiento de información',
         '- Guardan la información rápidamente y en cualquier orden.',
         '- Guardan de manera secuencial y por bloques; si se interrumpe, se pierden.',
         '- Guardan información mediante la "memoria muscular" y la asociación de acciones.'),
        ('Períodos de inactividad',
         '- Miran algo fijamente, se ponen a dibujar o a leer.',
         '- Canturrean para sí mismos o buscan hablar con alguien.',
         '- Tienen que moverse (caminan, se balancean, no se quedan quietos).'),
        ('Comunicación',
         '- Se impacientan si escuchan mucho rato. Usan palabras como "ver, aspecto, claro".',
         '- Les encanta escuchar y hablar. Hacen descripciones largas. Usan palabras como "sonar, ruido".',
         '- Gesticulan muchísimo. Se acercan mucho al otro. Usan palabras como "tomar, sentir".'),
        ('Distracción',
         '- Se distraen cuando hay movimiento o desorden visual; el ruido no molesta tanto.',
         '- Se distraen fácilmente cuando hay ruido o conversaciones cruzadas.',
         '- Se distraen si la explicación es puramente teórica/auditiva y no los involucra.')
    ]
    
    for i, data in enumerate(rows_data, start=1):
        row_cells = table1.rows[i].cells
        row_cells[0].text = data[0]
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = data[1]
        row_cells[2].text = data[2]
        row_cells[3].text = data[3]
    
    # Apply format to table cells
    for row in table1.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = 0
                p.paragraph_format.line_spacing = 1.0

    doc.add_paragraph() # spacing

    doc.add_heading('Estrategias prácticas en el aula', level=2)
    table2 = doc.add_table(rows=2, cols=5)
    table2.style = 'Table Grid'
    # Header
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Ámbito'
    hdr_cells2[1].text = 'Visual'
    hdr_cells2[2].text = 'Aural (Auditivo)'
    hdr_cells2[3].text = 'Lector/Escritor'
    hdr_cells2[4].text = 'Kinestésico'
    for cell in hdr_cells2:
        cell.paragraphs[0].runs[0].bold = True

    # Row 1
    row1_cells2 = table2.rows[1].cells
    row1_cells2[0].text = 'Ideas para estudiar mejor'
    row1_cells2[1].text = '- Llenar los apuntes de diagramas, mapas y dibujos.\n- Usar muchos símbolos y colores.\n- Armar sus propios mapas mentales.'
    row1_cells2[2].text = '- Grabarse resumiendo un texto.\n- Juntarse a estudiar debatiendo los temas.\n- Explicarle la lección a otra persona.'
    row1_cells2[3].text = '- Hacer sus propios glosarios o resúmenes de texto.\n- Pasar apuntes en limpio leyendo lo más importante.\n- Leer artículos extra.'
    row1_cells2[4].text = '- Pasar mucho tiempo en espacios prácticos o laboratorios.\n- Jugar a resolver un caso real.\n- Armar algún proyecto concreto o manualidad.'

    # Apply format
    for row in table2.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = 0
                p.paragraph_format.line_spacing = 1.0
                
    doc.add_page_break()
    doc.add_heading('Anexos Gráficos: Flujos del Modelo VAK', level=1)
    
    # Gráfico 1
    doc.add_heading('Gráfico 1: Flujo del Procesamiento de la Información según el Modelo VAK', level=2)
    doc.add_paragraph("Este diagrama muestra cómo la información del entorno es filtrada por nuestras preferencias sensoriales para convertirse en aprendizaje.")
    
    mermaid_1 = """graph TD
    A[Información del Entorno] --> B{Filtro de Preferencia Sensorial}
    B -->|Veo y ordeno| C[Canal VISUAL]
    B -->|Escucho y debato| D[Canal AUDITIVO]
    B -->|Toco y experimento| E[Canal KINESTÉSICO]
    C --> F[Procesamiento Cognitivo]
    D --> F
    E --> F
    F --> G((Aprendizaje y Retención))
    style A fill:#f9f,stroke:#333,stroke-width:2px
    style G fill:#bbf,stroke:#333,stroke-width:2px"""
    
    img1_path = 'c:\\Users\\HP 250 G10\\Documents\\GITHUT\\MonografiaUMA\\grafico1.png'
    if download_mermaid_image(mermaid_1, img1_path):
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.paragraph_format.first_line_indent = 0
        r_img1 = p_img1.add_run()
        r_img1.add_picture(img1_path, width=Inches(5.5))
    
    # Gráfico 2
    doc.add_heading('Gráfico 2: Ciclo de la Enseñanza Multimodal en el Aula', level=2)
    doc.add_paragraph("Este diagrama ilustra cómo un docente puede integrar los tres canales en una misma secuencia didáctica para asegurar que todos los alumnos se enganchen.")
    
    mermaid_2 = """flowchart LR
    1(Presentación del Tema) --> 2[Explicación Oral AUDITIVO]
    2 --> 3[Apoyo Gráfico/Pizarrón VISUAL]
    3 --> 4[Práctica o Simulación KINESTÉSICO]
    4 --> 5(((Consolidación del Aprendizaje)))
    style 1 fill:#ff9,stroke:#333,stroke-width:2px
    style 5 fill:#bfb,stroke:#333,stroke-width:2px"""
    
    img2_path = 'c:\\Users\\HP 250 G10\\Documents\\GITHUT\\MonografiaUMA\\grafico2.png'
    if download_mermaid_image(mermaid_2, img2_path):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.paragraph_format.first_line_indent = 0
        r_img2 = p_img2.add_run()
        r_img2.add_picture(img2_path, width=Inches(5.5))
    
    doc.save('c:\\Users\\HP 250 G10\\Documents\\GITHUT\\MonografiaUMA\\Monografia_VAK.docx')

if __name__ == '__main__':
    main()
